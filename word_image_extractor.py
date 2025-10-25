#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档图片提取工具
功能：从Word文档的表格中提取图片，并按照指定格式保存
"""

import os
import sys
import tkinter as tk
from tkinter import filedialog, simpledialog, scrolledtext, messagebox
from docx import Document
from PIL import Image
import io
import re
from docx.shared import Inches


def select_file(title="选择Word文档"):
    """让用户选择文件"""
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口
    file_path = filedialog.askopenfilename(
        title=title,
        filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
    )
    root.destroy()
    return file_path


def select_directory(title="选择输出目录"):
    """让用户选择目录"""
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口
    dir_path = filedialog.askdirectory(title=title)
    root.destroy()
    return dir_path


def sanitize_filename(filename):
    """清理文件名，移除或替换非法字符"""
    # 移除或替换Windows文件名中的非法字符
    invalid_chars = '<>"/\\|?*:'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    # 移除多余的空格和换行符
    filename = ' '.join(filename.split())
    # 限制文件名长度
    if len(filename) > 100:
        filename = filename[:100]
    return filename


def show_table_content(table):
    """显示表格内容，并让用户选择一个单元格作为Fname"""
    # 使用更简单的方式，避免创建两个Tk窗口
    print("正在准备表格预览...")
    
    # 预计算表格内容
    content = []
    try:
        # 大幅减少处理的行数和列数，确保快速响应
        max_rows = 10
        max_cols = 5
        rows_to_process = table.rows[:max_rows]
        
        for i, row in enumerate(rows_to_process):
            if i > 0:  # 跳过表头行外的其他行
                continue
            
            row_text = []
            cells_to_process = row.cells[:max_cols]
            
            for j, cell in enumerate(cells_to_process):
                cell_text = cell.text.strip()[:50]
                row_text.append(f"[{i},{j}]: {cell_text}")
            content.append(" | ".join(row_text))
        
        # 添加简化版提示
        content.append("\n\n[仅显示第一行作为参考，请输入单元格坐标，如：0,0]")
    except Exception as e:
        content.append(f"解析表格时出错: {str(e)}")
        print(f"表格解析错误: {str(e)}")
    
    print("表格预览内容准备完成")
    
    # 创建主窗口
    root = tk.Tk()
    root.title("选择Fname单元格")
    root.geometry("800x600")
    
    # 创建文本区域显示表格内容
    text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=90, height=30, font=("SimHei", 10))
    text_area.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
    
    # 显示表格内容
    text_area.insert(tk.END, "\n\n".join(content))
    text_area.config(state=tk.DISABLED)
    
    # 存储用户选择的单元格坐标
    selected_cell = [None, None]
    selected_text = [""]
    
    def on_select():
        try:
            row = int(row_entry.get())
            col = int(col_entry.get())
            if 0 <= row < len(table.rows) and 0 <= col < len(table.rows[row].cells):
                selected_cell[0] = row
                selected_cell[1] = col
                selected_text[0] = table.rows[row].cells[col].text.strip()
                root.destroy()
            else:
                messagebox.showerror("错误", "单元格坐标超出范围！")
        except ValueError:
            messagebox.showerror("错误", "请输入有效的数字！")
    
    # 创建输入框和按钮
    frame = tk.Frame(root)
    frame.pack(pady=10, padx=10, fill=tk.X)
    
    tk.Label(frame, text="请输入要选择的单元格坐标：").pack(side=tk.LEFT, padx=5)
    tk.Label(frame, text="行：").pack(side=tk.LEFT, padx=5)
    row_entry = tk.Entry(frame, width=5)
    row_entry.pack(side=tk.LEFT, padx=5)
    
    tk.Label(frame, text="列：").pack(side=tk.LEFT, padx=5)
    col_entry = tk.Entry(frame, width=5)
    col_entry.pack(side=tk.LEFT, padx=5)
    
    tk.Button(frame, text="确定", command=on_select).pack(side=tk.LEFT, padx=10)
    
    root.mainloop()
    
    return selected_text[0]


def extract_images_from_item(item, output_dir, base_name, item_index):
    """从单个item中提取图片并保存"""
    image_count = 0
    
    try:
        # 为每个item创建子文件夹
        item_folder = os.path.join(output_dir, f"{base_name}_{item_index}")
        os.makedirs(item_folder, exist_ok=True)
        
        # 尝试从表格中提取图片
        # 首先尝试直接从文档关系中获取所有图片
        try:
            for rel in item.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_count += 1
                        image_part = rel.target_part
                        image_bytes = image_part._blob
                        
                        # 确定图片格式
                        content_type = image_part.content_type
                        if 'png' in content_type:
                            ext = 'png'
                        elif 'jpeg' in content_type:
                            ext = 'jpg'
                        elif 'gif' in content_type:
                            ext = 'gif'
                        else:
                            ext = 'png'  # 默认使用png
                        
                        # 保存图片
                        image_path = os.path.join(item_folder, f"{base_name}_{image_count}.{ext}")
                        with open(image_path, 'wb') as f:
                            f.write(image_bytes)
                        
                        print(f"已保存图片: {image_path}")
                    except Exception as inner_e:
                        print(f"处理单个图片时出错: {inner_e}")
        except Exception as e:
            print(f"直接获取图片失败，尝试备用方法: {e}")
            
            # 备用方法：通过XPath查找
            try:
                for row in item.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # 添加命名空间处理
                                namespaces = {
                                    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                                }
                                
                                for shape in run.element.xpath('.//pic:pic', namespaces=namespaces):
                                    try:
                                        # 提取图片数据
                                        blip = shape.xpath('.//a:blip', namespaces=namespaces)[0]
                                        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                        
                                        # 获取图片
                                        if rId and rId in item.part.rels:
                                            image_count += 1
                                            image_part = item.part.rels[rId].target_part
                                            image_bytes = image_part._blob
                                            
                                            # 确定图片格式
                                            content_type = image_part.content_type
                                            if 'png' in content_type:
                                                ext = 'png'
                                            elif 'jpeg' in content_type:
                                                ext = 'jpg'
                                            elif 'gif' in content_type:
                                                ext = 'gif'
                                            else:
                                                ext = 'png'  # 默认使用png
                                            
                                            # 保存图片
                                            image_path = os.path.join(item_folder, f"{base_name}_{image_count}.{ext}")
                                            with open(image_path, 'wb') as f:
                                                f.write(image_bytes)
                                            
                                            print(f"已保存图片: {image_path}")
                                    except Exception as inner_e:
                                        print(f"提取图片时出错: {inner_e}")
            except Exception as e:
                print(f"备用方法也失败: {e}")
    except Exception as e:
        print(f"处理item时发生错误: {e}")
    
    return image_count


def main():
    """主函数"""
    print("===== Word文档图片提取工具 =====")
    
    # 使用命令行输入作为备选，避免GUI可能的问题
    print("提示：如果GUI操作卡住，可以在终端看到处理进度")
    
    try:
        # 让用户选择输入文件
        doc_path = select_file()
        if not doc_path:
            print("未选择文件，程序退出。")
            return
        
        print(f"已选择文件: {doc_path}")
        
        # 让用户选择输出目录
        output_dir = select_directory()
        if not output_dir:
            print("未选择输出目录，程序退出。")
            return
        
        print(f"已选择输出目录: {output_dir}")
        
        # 加载文档前的提示
        print("正在加载文档，请稍候...")
        
        # 加载文档
        try:
            doc = Document(doc_path)
            print(f"成功加载文档: {doc_path}")
        except Exception as e:
            print(f"加载文档失败: {str(e)}")
            messagebox.showerror("错误", f"加载文档失败: {str(e)}")
            return
        
        # 获取所有表格作为items
        items = doc.tables
        total_items = len(items)
        
        if total_items == 0:
            print("文档中没有找到表格！")
            messagebox.showinfo("提示", "文档中没有找到表格！")
            return
        
        print(f"在文档中找到 {total_items} 个表格（item）")
        
        # 简化Fname选择过程，使用命令行输入
        print("\n请输入要作为Fname的单元格坐标（如：0,0）:")
        print("例如，0,0 表示第一行第一列的单元格")
        
        # 创建一个简单的输入窗口
        root = tk.Tk()
        root.title("输入单元格坐标")
        root.geometry("400x200")
        
        label = tk.Label(root, text="请输入单元格坐标 (格式: 行号,列号，如0,0):")
        label.pack(pady=10)
        
        entry = tk.Entry(root, width=20)
        entry.pack(pady=10)
        entry.insert(0, "0,0")  # 默认值
        
        # 用于存储用户输入的变量
        result = [None]
        
        def on_submit():
            result[0] = entry.get()
            root.destroy()
        
        submit_btn = tk.Button(root, text="确定", command=on_submit)
        submit_btn.pack(pady=10)
        
        # 设置窗口关闭时的默认值
        root.protocol("WM_DELETE_WINDOW", lambda: (result.__setitem__(0, "0,0"), root.destroy()))
        
        # 确保窗口在最前面
        root.attributes('-topmost', True)
        root.mainloop()
        
        # 获取用户输入
        cell_coords = result[0]
        if not cell_coords:
            cell_coords = "0,0"
        
        try:
            # 解析坐标
            row_idx, col_idx = map(int, cell_coords.split(","))
            # 验证坐标是否有效
            if row_idx < 0 or col_idx < 0 or row_idx >= len(items[0].rows) or col_idx >= len(items[0].rows[row_idx].cells):
                print("坐标无效，使用默认值 (0,0)")
                row_idx, col_idx = 0, 0
            
            # 获取Fname值
            fname = items[0].rows[row_idx].cells[col_idx].text.strip()
            print(f"已选择Fname: {fname}")
        except Exception as e:
            print(f"解析坐标失败: {str(e)}，使用默认名称")
            fname = "默认名称"
        
        if not fname:
            print("未选择Fname，使用默认名称。")
            fname = "默认名称"
        
        # 清理Fname以用作文件名
        sanitized_fname = sanitize_filename(fname)
        print(f"已选择Fname: {fname}\n清理后的文件名: {sanitized_fname}")
        
        # 处理每个item
        total_images = 0
        created_folders = 0
        
        print("\n开始提取图片，按Ctrl+C可以随时中断...")
        
        for i, item in enumerate(items):
            try:
                print(f"\n处理第 {i+1}/{total_items} 个表格...")
                
                # 提取Fname值
                try:
                    # 使用之前获取的row_idx和col_idx
                    if row_idx < len(item.rows) and col_idx < len(item.rows[row_idx].cells):
                        item_fname = item.rows[row_idx].cells[col_idx].text.strip()
                        # 清理文件名
                        item_fname = sanitize_filename(item_fname)
                        if not item_fname:  # 如果清理后为空，使用默认名称
                            item_fname = f"item_{i+1}"
                    else:
                        print("坐标无效，使用默认名称")
                        item_fname = f"item_{i+1}"
                except Exception as e:
                    print(f"获取Fname失败: {e}")
                    item_fname = f"item_{i+1}"
                
                print(f"使用Fname: {item_fname}")
                
                # 提取图片 - 设置超时保护
                try:
                    import threading
                    import queue
                    
                    # 创建队列用于获取结果
                    result_queue = queue.Queue()
                    
                    # 定义线程函数
                    def extract_thread():
                        try:
                            count = extract_images_from_item(item, output_dir, item_fname, i+1)
                            result_queue.put(count)
                        except Exception as e:
                            print(f"线程中出错: {e}")
                            result_queue.put(0)
                    
                    # 创建并启动线程
                    thread = threading.Thread(target=extract_thread)
                    thread.daemon = True  # 设置为守护线程，主程序结束时自动终止
                    thread.start()
                    
                    # 等待线程完成，但设置超时时间为30秒
                    thread.join(timeout=30)
                    
                    # 检查线程是否还在运行
                    if thread.is_alive():
                        print(f"处理表格 {i+1} 超时，跳过")
                        images_count = 0
                    else:
                        # 从队列获取结果
                        images_count = result_queue.get()
                        
                    total_images += images_count
                    
                    if images_count > 0:
                        created_folders += 1
                    
                    print(f"该表格提取了 {images_count} 张图片")
                    
                except Exception as e:
                    print(f"图片提取过程出错: {e}")
                    continue
                
            except Exception as e:
                print(f"处理表格 {i+1} 时出错: {e}")
                continue
        
        # 显示处理结果
        print("\n===== 处理完成 =====")
        print(f"总计处理了 {total_items} 个表格")
        print(f"总计提取了 {total_images} 张图片")
        print(f"总计创建了 {created_folders} 个文件夹")
        print(f"所有图片已保存到目录: {output_dir}")
        
        # 显示完成消息框
        root = tk.Tk()
        root.withdraw()
        messagebox.showinfo("完成", f"处理完成！\n总计处理了 {total_items} 个表格\n总计提取了 {total_images} 张图片\n总计创建了 {created_folders} 个文件夹")
        root.destroy()
        
    except Exception as e:
        print(f"处理文档时出错: {e}")
        # 显示错误消息框
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror("错误", f"处理文档时出错: {str(e)}")
        root.destroy()


if __name__ == "__main__":
    # 检查是否安装了必要的库
    try:
        import docx
        from PIL import Image
    except ImportError:
        print("正在安装必要的库...")
        os.system(f"{sys.executable} -m pip install python-docx pillow")
        print("库安装完成，请重新运行程序。")
        input("按Enter键退出...")
        sys.exit(0)
    
    main()