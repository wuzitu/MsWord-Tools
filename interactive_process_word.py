import os
import re
import tkinter as tk
from tkinter import filedialog
from docx import Document

# --- 1. 辅助函数 ---

def sanitize_filename(name):
    """
    (需求 8) 清理文件名，去除Windows文件名中的非法字符，并去除首尾空格。
    """
    if not name:
        return "Untitled"
    # 去除非法字符 \ / : * ? " < > |
    name = re.sub(r'[\\/*?:"<>|]', '-', name)
    # 去除首尾的空格和换行符
    name = name.strip()
    # 如果清理后为空，则提供一个默认名
    if not name:
        return "Untitled"
    return name

def get_table_text_for_display(table):
    """
    (需求 2) 提取表格所有文本内容，格式化后用于在控制台显示。
    """
    text_output = []
    try:
        for r_idx, row in enumerate(table.rows):
            row_text = []
            for c_idx, cell in enumerate(row.cells):
                # 格式化输出，例如: [行0,列0]: 单元格内容
                cell_content = cell.text.strip().replace("\n", " ") # 将单元格内换行替换为空格
                row_text.append(f"[{r_idx},{c_idx}]: {cell_content}")
            
            # 用 " | " 分隔同一行的单元格
            text_output.append(" | ".join(row_text))
        
        # 用换行符分隔每一行
        return "\n".join(text_output)
    
    except Exception as e:
        return f"读取表格内容时出错: {e}"

# --- 2. 核心处理函数 ---

def process_document_interactive(doc_path, output_dir):
    """
    主处理逻辑：
    1. (需求 1) 遍历所有表格 (item)
    2. (需求 2) 显示item内容，等待用户输入Fname
    3. (需求 3) 创建Fname同名文件夹
    4. (需求 4) 提取该表格内的所有图片，并以Fname_序号命名
    """
    
    print(f"--- 开始处理文件: {doc_path} ---")
    
    try:
        document = Document(doc_path)
        tables = document.tables
        total_tables = len(tables)
        total_images_processed = 0
        
        if total_tables == 0:
            print("警告: 在此文档中未找到任何表格。程序退出。")
            return 0

        print(f"文档中总计 {total_tables} 个表格 (item)。")

        # 定义XML命名空间，用于查找图片
        ns_map = {
            'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
            'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        }

        # (需求 1) 遍历所有表格 (item)
        for i, table in enumerate(tables):
            print("\n" + "="*50)
            print(f"--- 正在处理表格 {i + 1}/{total_tables} ---")
            print("="*50)
            
            # (需求 2) 输出单个item的内容
            print("\n[表格内容预览]:")
            item_content = get_table_text_for_display(table)
            print(item_content)
            print("-"*50)

            # (需求 2 & 8) 让用户选择文字内容，作为变量Fname
            print("请从上面的内容中，复制您想用作文件夹/图片名称的文本，")
            fname_raw = input("然后粘贴到这里 (或手动输入) 并按 Enter 键: ")
            
            # (需求 3) 定义文件夹名称
            Fname = sanitize_filename(fname_raw)
            if not Fname:
                Fname = f"Item_{i + 1}_Untitled"
                print(f"  输入为空，使用默认名称: {Fname}")
            else:
                print(f"  已获取 Fname: '{Fname}'")

            # 创建文件夹
            target_folder_path = os.path.join(output_dir, Fname)
            os.makedirs(target_folder_path, exist_ok=True)
            print(f"  已创建/确认文件夹: {target_folder_path}")

            # (需求 4) 提取图片并重命名
            image_counter = 0
            # 遍历表格的 -> 行 -> 单元格 -> 段落 -> 运行(run)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            # 查找run中的 'blip' (Bitmap Location and Identification) 元素
                            try:
                                blip_list = run.element.xpath('.//a:blip', namespaces=ns_map)
                                if blip_list:
                                    # 获取图片的 rId (relationship Id)
                                    rId = blip_list[0].get(f'{{{ns_map["r"]}}}embed')
                                    if rId:
                                        # 通过rId从文档中获取图片部件
                                        image_part = document.part.related_parts[rId]
                                        image_blob = image_part.blob
                                        # 获取图片扩展名
                                        image_ext = image_part.default_part_name.split('.')[-1]

                                        # (需求 5) 定义图片文件名
                                        image_counter += 1
                                        total_images_processed += 1
                                        image_name = f"{Fname}_{image_counter}.{image_ext}"
                                        image_save_path = os.path.join(target_folder_path, image_name)
                                        
                                        # 保存图片
                                        with open(image_save_path, 'wb') as f:
                                            f.write(image_blob)
                            except Exception as e:
                                print(f"  提取图片时出错: {e}")

            if image_counter == 0:
                print(f"  在 Fname: '{Fname}' 的表格中未找到图片。")
            else:
                print(f"  成功提取 {image_counter} 张图片。")

        # (需求 5) 处理完成后，显示结果
        print("\n" + "="*50)
        print("--- 所有任务处理完毕 ---")
        print(f"总计 {total_tables} 个表格 (item) 已处理完毕。")
        print(f"总计提取 {total_images_processed} 张图片。")
        print("="*50)
        return total_tables

    except Exception as e:
        print(f"\n--- 发生严重错误 ---")
        print(f"处理文件失败: {e}")
        print("请确保文件未被打开，且具有读取权限。")
        return 0

# --- 3. 主程序入口 ---
def main():
    # (需求 6) 最好能让用户选择输入的word文档、输出的文件夹目录
    # 弹出GUI窗口让用户选择
    root = tk.Tk()
    root.withdraw() # 隐藏主窗口

    print("请在弹出的窗口中，选择您要处理的 Word 文档 (.docx)...")
    doc_path = filedialog.askopenfilename(
        title="请选择一个 Word 文档",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not doc_path:
        print("用户取消了选择。程序退出。")
        return

    print("请在弹出的窗口中，选择图片要导出到的目标文件夹...")
    output_dir = filedialog.askdirectory(
        title="请选择一个输出文件夹"
    )
    if not output_dir:
        print("用户取消了选择。程序退出。")
        return

    # 调用核心处理函数
    process_document_interactive(doc_path, output_dir)
    
    # 防止exe窗口闪退
    print("\n按 Enter 键退出...")
    input()

# ---------------------------------
if __name__ == "__main__":
    main()
# ---------------------------------