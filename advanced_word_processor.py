import os
import re
import tkinter as tk
from tkinter import filedialog
from docx import Document
from datetime import datetime
from lxml.etree import QName # 用于兼容地处理 XML 命名空间

# --- 1. 配置 & 日志变量 ---

# 错误日志列表，最多记录500条
ERROR_LOGS = []
MAX_LOG_ENTRIES = 500

# 统计变量
TOTAL_TABLES = 0
PROCESSED_FOLDERS = 0
TOTAL_IMAGES = 0

# --- 2. 辅助函数 ---

def log_error(message):
    """记录错误日志"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_entry = f"[{timestamp}] {message}"
    if len(ERROR_LOGS) >= MAX_LOG_ENTRIES:
        ERROR_LOGS.pop(0) # 移除最旧的
    ERROR_LOGS.append(log_entry)
    print(f"[错误记录]: {message}")

def sanitize_filename(name):
    """清理文件名，去除Windows文件名中的非法字符"""
    if not name:
        return "Untitled"
    name = re.sub(r'[\\/*?:"<>|]', '-', name)
    name = name.strip()
    return name if name else "Untitled"

def parse_cell_index(user_input):
    """
    解析用户输入的单元格编号，格式如 "0,0" (row_index,col_index)
    """
    if not user_input or ',' not in user_input:
        return None
    try:
        r, c = map(str.strip, user_input.split(','))
        r_idx = int(r)
        c_idx = int(c)
        if r_idx < 0 or c_idx < 0:
            return None # 索引不能为负
        return r_idx, c_idx
    except ValueError:
        return None # 无法转换为整数

# --- 3. 核心处理函数 ---

def process_document(doc_path, output_dir, target_cell):
    """
    主处理逻辑：自动根据target_cell从每个表格中提取Fname
    """
    global PROCESSED_FOLDERS, TOTAL_IMAGES, TOTAL_TABLES
    
    print(f"--- 开始处理文件: {doc_path} ---")
    
    try:
        document = Document(doc_path)
        tables = document.tables
        TOTAL_TABLES = len(tables)
        
        if TOTAL_TABLES == 0:
            print("警告: 在此文档中未找到任何表格。")
            return

        print(f"文档中总计 {TOTAL_TABLES} 个表格 (item)。")

        # 定义需要查找的 XML 元素的完全限定名，解决 BaseOxmlElement.xpath() 错误
        r_embed_qname = QName("http://schemas.openxmlformats.org/officeDocument/2006/relationships", 'embed')
        a_blip_qname = QName("http://schemas.openxmlformats.org/drawingml/2006/main", 'blip')

        row_idx, col_idx = target_cell # 固定的目标单元格索引

        # 遍历所有表格 (item)
        for i, table in enumerate(tables):
            print(f"\n--- 正在处理表格 {i + 1}/{TOTAL_TABLES} ---")
            
            # --- 自动获取 Fname (统一单元格逻辑) ---
            Fname = f"Item_{i+1}_Untitled"
            try:
                # 尝试获取用户指定的单元格内容作为Fname
                fname_raw = table.cell(row_idx, col_idx).text
                Fname = sanitize_filename(fname_raw)
                
                if not Fname:
                    Fname = f"Item_{i+1}_Untitled"
                    log_error(f"表格 {i+1}: 目标单元格 ({row_idx},{col_idx}) 内容为空或仅含非法字符，使用默认命名。")
                print(f"  单元格 ({row_idx},{col_idx}) Fname: '{Fname}'")
                
            except IndexError:
                log_error(f"表格 {i+1}: 目标单元格 ({row_idx},{col_idx}) 不存在，跳过此表格。")
                print(f"  目标单元格 ({row_idx},{col_idx}) 不存在，跳过。")
                continue
            except Exception as e:
                log_error(f"表格 {i+1}: 获取 Fname 时发生未知错误: {e}")
                continue

            # --- 创建文件夹 ---
            target_folder_path = os.path.join(output_dir, Fname)
            try:
                os.makedirs(target_folder_path, exist_ok=True)
                PROCESSED_FOLDERS += 1
                print(f"  已创建/确认文件夹: {target_folder_path}")
            except Exception as e:
                log_error(f"表格 {i+1}: 创建文件夹失败 ({target_folder_path}): {e}")
                continue

            # --- 提取图片 ---
            image_counter = 0
            
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            try:
                                # 使用 findall 和 QName 替换 xpath，解决兼容性问题
                                blip_list = run.element.findall('.//' + str(a_blip_qname))
                                
                                if blip_list:
                                    # 获取 r:embed 属性值
                                    rId = blip_list[0].get(r_embed_qname)
                                    
                                    if rId:
                                        # 通过 rId 从文档中获取图片部件
                                        image_part = document.part.related_parts[rId]
                                        image_blob = image_part.blob
                                        image_ext = image_part.default_part_name.split('.')[-1]

                                        # 定义图片文件名 (Fname + 数字序号)
                                        image_counter += 1
                                        TOTAL_IMAGES += 1
                                        image_name = f"{Fname}_{image_counter}.{image_ext}"
                                        image_save_path = os.path.join(target_folder_path, image_name)
                                        
                                        with open(image_save_path, 'wb') as f:
                                            f.write(image_blob)
                            except Exception as e:
                                # 记录提取图片时的任何错误
                                log_error(f"表格 {i+1}, Fname '{Fname}': 提取或保存图片时出错: {e}")

            if image_counter == 0:
                print(f"  在 '{Fname}' 的表格中未找到图片。")
            else:
                print(f"  成功提取 {image_counter} 张图片。")


    except Exception as e:
        log_error(f"处理文档时发生致命错误: {e}")
        print(f"\n--- 发生致命错误 ---")
        print(f"处理文件失败: {e}")

# --- 4. 主程序入口 ---
def main():
    global PROCESSED_FOLDERS, TOTAL_IMAGES
    
    # 隐藏Tkinter主窗口
    root = tk.Tk()
    root.withdraw() 
    
    # --- 步骤 1: 获取单元格编号 ---
    target_cell = None
    while target_cell is None:
        print("\n" + "="*50)
        print("请定义所有表格用于命名的单元格编号 (例如: 0,0 代表第一行第一列):")
        cell_input = input("输入行,列编号 (row_index,col_index): ")
        target_cell = parse_cell_index(cell_input)
        if target_cell is None:
            print("输入格式错误或索引无效。请重新输入。")

    # --- 步骤 2: 选择文件和目录 ---
    print("\n请在弹出的窗口中，选择您要处理的 Word 文档 (.docx)...")
    doc_path = filedialog.askopenfilename(
        title="请选择一个 Word 文档",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not doc_path:
        print("用户取消了文件选择。程序退出。")
        return

    print("请在弹出的窗口中，选择图片要导出到的目标文件夹...")
    output_dir = filedialog.askdirectory(
        title="请选择一个输出文件夹"
    )
    if not output_dir:
        print("用户取消了输出目录选择。程序退出。")
        return
        
    print(f"\n[配置]: 目标单元格为：第 {target_cell[0]+1} 行，第 {target_cell[1]+1} 列。")
    print(f"[注意]: 程序将全自动运行。")

    # --- 步骤 3: 调用核心处理函数 ---
    process_document(doc_path, output_dir, target_cell)
    
    # --- 步骤 4: 结果输出 ---
    
    # 保存日志
    log_file_path = os.path.join(output_dir, "error_log.txt")
    try:
        with open(log_file_path, 'w', encoding='utf-8') as f:
            f.write("--- 错误和警告日志记录 ---\n")
            f.write(f"文件处理时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            if ERROR_LOGS:
                f.write("\n".join(ERROR_LOGS))
            else:
                f.write("未记录到任何错误或警告。\n")
        print(f"\n[日志]: 错误日志已保存到: {log_file_path}")
    except Exception as e:
        print(f"[日志错误]: 无法保存日志文件: {e}")
    
    # 最终统计结果
    print("\n" + "="*50)
    print("--- 最终统计结果 ---")
    print(f"总计检测到表格数量: {TOTAL_TABLES}")
    print(f"成功创建的文件夹数量: {PROCESSED_FOLDERS}")
    print(f"提取的图片总数量: {TOTAL_IMAGES}")
    print(f"错误日志条数: {len(ERROR_LOGS)} / {MAX_LOG_ENTRIES}")
    print("========================")
    
    # 防止exe窗口闪退
    print("\n处理完成。按 Enter 键退出...")
    input()

# ---------------------------------
if __name__ == "__main__":
    main()
# ---------------------------------