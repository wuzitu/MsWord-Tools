import os
import hashlib
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.shared import Inches


def select_word_file():
    """
    选择 Word 文件
    """
    file_path = filedialog.askopenfilename(
        title="请选择Word文档",
        filetypes=[("Word文件", "*.docx")]
    )
    return file_path


def select_output_dir():
    """
    选择输出文件夹
    """
    folder_path = filedialog.askdirectory(
        title="请选择图片输出目录"
    )
    return folder_path


def save_table_as_text(table, output_path):
    """
    将表格内容保存为文本文件
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            f.write('\t'.join(row_data) + '\n')


def save_table_as_docx(table, output_path):
    """
    将表格内容保存为新的Word文档
    """
    new_doc = Document()
    new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    
    # 复制表格内容
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            new_table.cell(i, j).text = cell.text.strip()
    
    new_doc.save(output_path)


def extract_context_around_table(doc, table_index, output_path):
    """
    提取表格周围的上下文信息（包括前后表格和文本段落）
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(f"=== 表格 {table_index + 1} 的上下文信息 ===\n\n")
        
        # 获取文档中的所有元素
        all_elements = []
        for element in doc.element.body:
            all_elements.append(element)
        
        # 找到表格在文档中的位置
        table_element = doc.tables[table_index]._element
        table_position = -1
        for i, element in enumerate(all_elements):
            if element is table_element:
                table_position = i
                break
        
        if table_position == -1:
            f.write("无法定位表格在文档中的位置。\n")
            return
        
        # 提取表格前的内容（最多5个元素）
        f.write("【表格前的内容】\n")
        start_pos = max(0, table_position - 5)
        for i in range(start_pos, table_position):
            element = all_elements[i]
            if element.tag.endswith('tbl'):  # 表格
                # 尝试找到对应的python-docx表格对象
                for j, table in enumerate(doc.tables):
                    if table._element is element:
                        f.write(f"[前表格 {j+1}]\n")
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            f.write('\t'.join(row_data) + '\n')
                        f.write('\n')
                        break
            elif element.tag.endswith('p'):  # 段落
                paragraph_text = element.text.strip() if element.text else ""
                if paragraph_text:
                    f.write(f"[前段落] {paragraph_text}\n")
        f.write("\n")
        
        # 提取当前表格
        f.write("【当前表格内容】\n")
        current_table = doc.tables[table_index]
        for row in current_table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            f.write('\t'.join(row_data) + '\n')
        f.write("\n")
        
        # 提取表格后的内容（最多5个元素）
        f.write("【表格后的内容】\n")
        end_pos = min(len(all_elements), table_position + 6)  # +6 because range is exclusive
        for i in range(table_position + 1, end_pos):
            element = all_elements[i]
            if element.tag.endswith('tbl'):  # 表格
                # 尝试找到对应的python-docx表格对象
                for j, table in enumerate(doc.tables):
                    if table._element is element:
                        f.write(f"[后表格 {j+1}]\n")
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            f.write('\t'.join(row_data) + '\n')
                        f.write('\n')
                        break
            elif element.tag.endswith('p'):  # 段落
                paragraph_text = element.text.strip() if element.text else ""
                if paragraph_text:
                    f.write(f"[后段落] {paragraph_text}\n")


def extract_images_from_cell(cell, output_folder, fname_base, image_counter, seen_hashes):
    """
    提取单元格中的图片（兼容旧版 python-docx，无 namespaces 参数）
    返回提取的图片数量
    """
    count = 0
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # 查找所有 blip 节点（嵌入图片）
            inline_shapes = run.element.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
            )
            for blip in inline_shapes:
                rid = blip.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if rid:
                    image_part = run.part.related_parts[rid]
                    img_bytes = image_part.blob
                    
                    # 计算图片的哈希值用于去重
                    img_hash = hashlib.md5(img_bytes).hexdigest()
                    
                    # 如果图片已经处理过，则跳过
                    if img_hash in seen_hashes:
                        continue
                    
                    # 标记图片已处理
                    seen_hashes.add(img_hash)
                    count += 1
                    # 修改图片命名方式，使用横杠分隔
                    img_filename = f"{fname_base}-{image_counter + count}.png"
                    img_path = os.path.join(output_folder, img_filename)
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
    return count


class WordImageExtractorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Word表格图片提取工具 - H3C AI助手")

        self.word_file = ""
        self.output_dir = ""
        self.doc = None
        self.tables = []
        self.coord = None  # 修改为None，表示尚未选择
        self.fname_first = ""

        # 按钮选择文件和目录（按顺序）
        self.file_btn = tk.Button(root, text="1. 选择Word文档", command=self.load_word_file)
        self.file_btn.pack(pady=5)
        
        self.dir_btn = tk.Button(root, text="2. 选择输出目录", command=self.load_output_dir, state="disabled")
        self.dir_btn.pack(pady=5)
        
        self.coord_label = tk.Label(root, text="3. 请先选择Word文档和输出目录")
        self.coord_label.pack(pady=5)

        # 表格显示框
        self.table_frame = tk.Frame(root)
        self.table_frame.pack(pady=10)

        # 进度条
        self.progress = ttk.Progressbar(root, orient="horizontal", length=300, mode="determinate")
        self.progress.pack(pady=5)

        # 开始按钮
        self.process_btn = tk.Button(root, text="4. 开始处理", command=self.process_tables, state="disabled")
        self.process_btn.pack(pady=10)

    def load_word_file(self):
        """
        加载 Word 文件并显示第一个表格
        """
        self.word_file = select_word_file()
        if not self.word_file:
            return
        
        try:
            self.doc = Document(self.word_file)
            self.tables = self.doc.tables
            if not self.tables:
                messagebox.showerror("错误", "文档中没有表格")
                return
            
            messagebox.showinfo("提示", f"已加载 {len(self.tables)} 个表格")
            # 启用输出目录选择按钮
            self.dir_btn.config(state="normal")
            self.coord_label.config(text="2. 请选择输出目录，然后选择Fname坐标")
        except Exception as e:
            messagebox.showerror("错误", f"加载文档失败: {str(e)}")
            return

    def load_output_dir(self):
        """
        选择输出目录
        """
        if not self.word_file:
            messagebox.showerror("错误", "请先选择Word文档")
            return
            
        self.output_dir = select_output_dir()
        if self.output_dir:
            messagebox.showinfo("提示", f"输出目录已选择：{self.output_dir}")
            # 显示第一个表格供选择坐标
            self.display_first_table()
            self.coord_label.config(text="3. 请点击表格中的单元格选择Fname坐标")

    def display_first_table(self):
        """
        显示第一个表格的每个单元格按钮供选择
        """
        # 只有当文档和输出目录都选择了才显示表格
        if not self.word_file or not self.output_dir:
            return
            
        for widget in self.table_frame.winfo_children():
            widget.destroy()

        first_table = self.tables[0]
        tk.Label(self.table_frame, text="请选择作为Fname的单元格:").pack()
        for r, row in enumerate(first_table.rows):
            row_frame = tk.Frame(self.table_frame)
            row_frame.pack()
            for c, cell in enumerate(row.cells):
                text_show = cell.text.strip()[:15] or "[空]"
                btn = tk.Button(
                    row_frame,
                    text=text_show,
                    command=lambda rr=r, cc=c: self.set_coord(rr, cc)
                )
                btn.pack(side="left", padx=2, pady=2)

    def set_coord(self, r, c):
        """
        选择 Fname 坐标
        """
        if not self.word_file or not self.output_dir:
            messagebox.showerror("错误", "请先选择Word文档和输出目录")
            return
            
        self.coord = (r, c)
        self.fname_first = self.tables[0].cell(r, c).text.strip() or "未命名文件夹1"
        messagebox.showinfo("选择坐标", f"已选择坐标: ({r},{c})\nFname: {self.fname_first}")
        # 启用开始处理按钮
        self.process_btn.config(state="normal")
        self.coord_label.config(text=f"已选择坐标: ({r},{c}) Fname: {self.fname_first}")

    def process_tables(self):
        """
        按照坐标提取所有表格的图片
        """
        if not self.word_file:
            messagebox.showerror("错误", "请先选择Word文档")
            return
        if not self.output_dir:
            messagebox.showerror("错误", "请先选择输出目录")
            return
        if not self.coord:
            messagebox.showerror("错误", "请先选择Fname坐标")
            return

        row_idx, col_idx = self.coord
        total_tables = len(self.tables)
        total_images = 0  # 总图片计数
        total_unique_images = 0  # 唯一图片计数

        self.progress["maximum"] = total_tables
        self.progress["value"] = 0
        self.root.update_idletasks()

        for idx, table in enumerate(self.tables, start=1):
            # 确定文件夹名称
            try:
                fname_current = table.cell(row_idx, col_idx).text.strip() or f"未命名文件夹{idx}"
            except Exception:
                fname_current = f"未命名文件夹{idx}"

            item_folder = os.path.join(self.output_dir, fname_current)
            os.makedirs(item_folder, exist_ok=True)

            # 为每个表格创建独立的哈希集合，确保同一表格内的重复图片不会被提取
            seen_hashes = set()
            
            # 提取图片，使用全局计数器确保唯一性
            image_count = 0
            unique_image_count = 0
            table_image_counter = 0  # 每个表格的图片计数器
            for row in table.rows:
                for cell in row.cells:
                    # 传递当前表格的图片计数器和哈希集合
                    extracted = extract_images_from_cell(cell, item_folder, fname_current, table_image_counter, seen_hashes)
                    image_count += extracted
                    table_image_counter += extracted
                    unique_image_count += extracted

            total_images += image_count
            total_unique_images += unique_image_count
            
            # 如果文件夹名称是"未命名文件夹"开头，保存表格内容和上下文信息以便核对
            if fname_current.startswith("未命名文件夹"):
                # 保存为文本文件
                txt_path = os.path.join(item_folder, f"{fname_current}_表格内容.txt")
                save_table_as_text(table, txt_path)
                
                # 保存为Word文档
                docx_path = os.path.join(item_folder, f"{fname_current}_表格内容.docx")
                save_table_as_docx(table, docx_path)
                
                # 保存上下文信息
                context_path = os.path.join(item_folder, f"{fname_current}_上下文信息.txt")
                extract_context_around_table(self.doc, idx-1, context_path)
                
                print(f"[表格 {idx}] 提取图片 {unique_image_count} 张（共发现 {image_count} 张，去重后 {unique_image_count} 张），保存到 {item_folder}")
                print(f"        已保存表格内容到 {txt_path} 和 {docx_path}")
                print(f"        已保存上下文信息到 {context_path}")
            else:
                print(f"[表格 {idx}] 提取图片 {unique_image_count} 张（共发现 {image_count} 张，去重后 {unique_image_count} 张），保存到 {item_folder}")

            # 更新进度条
            self.progress["value"] = idx
            self.root.update_idletasks()

        messagebox.showinfo("处理完成", f"总计处理 {total_tables} 个表格，发现 {total_images} 张图片，去重后提取 {total_unique_images} 张唯一图片！")


if __name__ == "__main__":
    root = tk.Tk()
    app = WordImageExtractorGUI(root)
    root.mainloop()