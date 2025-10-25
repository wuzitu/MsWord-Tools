import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document


def select_word_file():
    """
    选择 Word 文件
    """
    file_path = filedialog.askopenfilename(
        title="请选择Word文档",
        filetypes=[("Word文件", "*.docx")]
    )
    return file_path


def select_output_dir():
    """
    选择输出文件夹
    """
    folder_path = filedialog.askdirectory(
        title="请选择图片输出目录"
    )
    return folder_path


def extract_images_from_cell(cell, output_folder, fname_base):
    """
    提取单元格中的图片（兼容旧版 python-docx，无 namespaces 参数）
    """
    count = 0
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # 查找所有 blip 节点（嵌入图片）
            inline_shapes = run.element.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
            )
            for blip in inline_shapes:
                rid = blip.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if rid:
                    image_part = run.part.related_parts[rid]
                    img_bytes = image_part.blob
                    count += 1
                    img_filename = f"{fname_base}{count}.png"
                    img_path = os.path.join(output_folder, img_filename)
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
    return count


class WordImageExtractorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Word表格图片提取工具 - H3C AI助手")

        self.word_file = ""
        self.output_dir = ""
        self.doc = None
        self.tables = []
        self.coord = (0, 0)
        self.fname_first = ""

        # 按钮选择文件和目录
        tk.Button(root, text="选择Word文档", command=self.load_word_file).pack(pady=5)
        tk.Button(root, text="选择输出目录", command=self.load_output_dir).pack(pady=5)

        # 表格显示框
        self.table_frame = tk.Frame(root)
        self.table_frame.pack(pady=10)

        # 进度条
        self.progress = ttk.Progressbar(root, orient="horizontal", length=300, mode="determinate")
        self.progress.pack(pady=5)

        # 开始按钮
        self.process_btn = tk.Button(root, text="开始处理", command=self.process_tables, state="disabled")
        self.process_btn.pack(pady=10)

    def load_word_file(self):
        """
        加载 Word 文件并显示第一个表格
        """
        self.word_file = select_word_file()
        if not self.word_file:
            return
        self.doc = Document(self.word_file)
        self.tables = self.doc.tables
        if not self.tables:
            messagebox.showerror("错误", "文档中没有表格")
            return
        messagebox.showinfo("提示", f"已加载 {len(self.tables)} 个表格，请选择Fname坐标")
        self.display_first_table()

    def load_output_dir(self):
        """
        选择输出目录
        """
        self.output_dir = select_output_dir()
        if self.output_dir:
            messagebox.showinfo("提示", f"输出目录已选择：{self.output_dir}")

    def display_first_table(self):
        """
        显示第一个表格的每个单元格按钮供选择
        """
        for widget in self.table_frame.winfo_children():
            widget.destroy()

        first_table = self.tables[0]
        for r, row in enumerate(first_table.rows):
            row_frame = tk.Frame(self.table_frame)
            row_frame.pack()
            for c, cell in enumerate(row.cells):
                text_show = cell.text.strip()[:15] or "[空]"
                btn = tk.Button(
                    row_frame,
                    text=text_show,
                    command=lambda rr=r, cc=c: self.set_coord(rr, cc)
                )
                btn.pack(side="left", padx=2, pady=2)

    def set_coord(self, r, c):
        """
        选择 Fname 坐标
        """
        self.coord = (r, c)
        self.fname_first = self.tables[0].cell(r, c).text.strip() or "未命名文件夹1"
        messagebox.showinfo("选择坐标", f"已选择坐标: ({r},{c})\nFname: {self.fname_first}")
        if self.word_file and self.output_dir:
            self.process_btn.config(state="normal")

    def process_tables(self):
        """
        按照坐标提取所有表格的图片
        """
        if not self.word_file or not self.output_dir:
            messagebox.showerror("错误", "请先选择Word文件和输出目录")
            return

        row_idx, col_idx = self.coord
        total_tables = len(self.tables)

        self.progress["maximum"] = total_tables
        self.progress["value"] = 0
        self.root.update_idletasks()

        for idx, table in enumerate(self.tables, start=1):
            # 确定文件夹名称
            try:
                fname_current = table.cell(row_idx, col_idx).text.strip() or f"未命名文件夹{idx}"
            except Exception:
                fname_current = f"未命名文件夹{idx}"

            item_folder = os.path.join(self.output_dir, fname_current)
            os.makedirs(item_folder, exist_ok=True)

            # 提取图片
            image_count = 0
            for row in table.rows:
                for cell in row.cells:
                    image_count += extract_images_from_cell(cell, item_folder, fname_current)

            print(f"[表格 {idx}] 提取图片 {image_count} 张，保存到 {item_folder}")

            # 更新进度条
            self.progress["value"] = idx
            self.root.update_idletasks()

        messagebox.showinfo("处理完成", f"总计 {total_tables} 个表格，已处理完毕！")


if __name__ == "__main__":
    root = tk.Tk()
    app = WordImageExtractorGUI(root)
    root.mainloop()